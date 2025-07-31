from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import random
import numpy as np

document = Document()

# document.add_heading('Document Title', 0)

SECTION_TEXTS = np.array([
    [   # chapter 1
        "1.1 Time Constant",
        "1.2 Electrostatic and Electromagnetic Fields",
        "1.3 Series Resonance",
        "1.4 Parallel Resonance",
        "1.5 Quality Factor",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
    ],
    [   # chapter 2
        "2.1 Semiconductors",
        "2.2 Diodes",
        "2.3 Transistors (BJT)",
        "2.4 Field Effect Transistors (FET)",
        "2.5 Silicon Controlled Rectifier (SCR)",
        "2.6 Amplifier Classes",
        "2.7 Amplifier Circuits",
        "2.8 Operational Amplifiers",
        "2.9 Mixers, Frequency Multipliers",
        "2.10 Digital Logic Elements",
        "2.11 Quartz Crystals",
        "2.12 Advanced Filter Circuits"
    ],
    [
        "3.1 AC Voltage Measurement",
        "3.2 Peak Envelope Power",
        "3.3 Dip Meters",
        "3.4 Frequency Counters",
        "3.5 Oscilloscopes",
        "3.6 Meters",
        "",
        "",
        "",
        "",
        "",
        ""
    ],
    [
        "4.1 Rectifiers",
        "4.2 Power Supply Filters",
        "4.3 Voltage Regulators",
        "4.4 Regulated Power Supplies",
        "",
        "",
        "",
        "",
        "",
        "",
        "",
        ""
    ],
    [
        "5.1 Oscillator Types",
        "5.2 Amplifier Essentials",
        "5.3 Transmitter Circuits",
        "5.4 Single Sideband",
        "5.5 FM Signals",
        "5.6 FM Transmitters",
        "5.7 Signal Processing",
        "5.8 Digital Radio",
        "5.9 Spread Spectrum",
        "",
        "",
        ""
    ],
    [
        "6.1 Receiver Essentials",
        "6.2 Receiver Dynamics",
        "6.3 Receiver Performance",
        "6.4 Receiver Circuitry",
        "6.5 Receiver Fundamentals",
        "",
        "",
        "",
        "",
        "",
        "",
        ""
    ],
    [
        "7.1 Antenna Tuning",
        "7.2 Impedance dynamics",
        "7.3 Antenna Matching",
        "7.4 Half Wave Dipoles",
        "7.5 Antenna Polarization",
        "7.6 Effective Radiated Power (ERP)",
        "7.7 Antenna Elevation",
        "7.8 Antenna Radiation Patterns",
        "7.9 Waveguides",
        "",
        "",
        ""
    ]
])
CHAPTER_TEXTS = np.array([
    "Chapter 1 - Advanced Theory",
    "Chapter 2 - Components and Circuits",
    "Chapter 3 - Electrical Measurement Techniques and Instruments",
    "Chapter 4 - Power Supplies",
    "Chapter 5 - Signal Processing",
    "Chapter 6 - Superheterodyne Receivers",
    "Chapter 7 - Antennas, Feedlines and Matching"
])

style = document.styles['Normal']
font = style.font

font.name = 'Arial'  # Change to any font you want
font.size = Pt(10)
def checkSectionOverflow(q_number):
    numbers = q_number.split('-')
    number = int(numbers[-1])
    section = int(numbers[-2])
    chapter = int(numbers[-3])
    # print(SECTION_TEXTS[chapter, section])
    if section == 1 and number == 1:
        # add new chapter thing
        p = document.add_paragraph()
        p.add_run(CHAPTER_TEXTS[chapter - 1]).bold = True
    if number == 1:
        # add new section thing

        p = document.add_paragraph()
        p.add_run(SECTION_TEXTS[chapter - 1, section - 1]).bold = True
    print(number)

with open("amat_adv_quest_delim.txt", 'r') as f:
    print(f.readline())
    # print(f.readline())
    for i in range(549):
        entries = f.readline().split(';')
        entries = entries[0:6]  # remove french lines

        q_number = entries[0]
        q_prompt = entries[1]
        q_correct = entries[2]
        # q_incor_1 = entries[3]
        # q_incor_2 = entries[4]
        # q_incor_3 = entries[5]

        checkSectionOverflow(q_number)

        choices = entries[2:6]
        letters = ['A', 'B', 'C', 'D']

        random.shuffle(choices)

        for index, item in enumerate(choices):
            if item == q_correct:
                correct_letter = letters[index]

        p = document.add_paragraph(q_number)
        p.add_run(f'        ({correct_letter})').bold = True

        p = document.add_paragraph(q_prompt)

        for letter, choice in zip(letters, choices):
            p = document.add_paragraph(style='List')  # Applies list formatting
            p.text = f"{letter}     {choice}"  # Override number with A, B, C, etc.
        print(entries)


document.add_page_break()

document.save('demo.docx')
